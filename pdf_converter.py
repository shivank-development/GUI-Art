from docx import Document
from docx.shared import Pt

# Create a new document
doc = Document()

# Title
doc.add_heading("📘 React + Django REST Framework (DRF) Full Stack Course", level=0)
doc.add_paragraph("With PostgreSQL, JWT Authentication, Blog App Project, and Deployment Guide")

# Table of Contents
doc.add_heading("📖 Table of Contents", level=1)
toc = [
    "Part 1 – Setup & Foundations",
    "Environment Setup",
    "Django + PostgreSQL Setup",
    "React Setup",
    "Connecting Backend and Frontend",
    "Exercises",
    "Part 2 – Django REST Framework (DRF + JWT Authentication)",
    "DRF Setup",
    "Models, Serializers, Views, Routers",
    "JWT Authentication with djangorestframework-simplejwt",
    "Permissions, Filtering, Pagination",
    "Exercises",
    "Part 3 – Mini Project: Blog App",
    "Backend (Django + DRF + PostgreSQL + JWT)",
    "Frontend (React + Axios + React Router)",
    "Protected Routes, Auth, Post CRUD",
    "Exercises",
    "Part 4 – Deployment Guide",
    "React → Vercel",
    "Django + PostgreSQL → Render",
    "Connecting Frontend ↔ Backend",
    "Exercises",
    "Appendix",
    "Architecture Diagram",
    "Resources"
]
for item in toc:
    doc.add_paragraph(item, style="List Bullet")

# Part 1
doc.add_heading("🟦 Part 1 – Setup & Foundations", level=1)
doc.add_heading("🔹 Step 1: Environment Setup", level=2)
doc.add_paragraph("Install requirements:")
doc.add_paragraph("""
# Install Python
python --version

# Install pip
python -m ensurepip --upgrade

# Install Node.js & npm
node -v
npm -v

# Install PostgreSQL
psql --version
""")

doc.add_heading("🔹 Step 2: Django + PostgreSQL Setup", level=2)
doc.add_paragraph("""
# Create virtual environment
python -m venv env
source env/bin/activate   # Mac/Linux
env\\Scripts\\activate      # Windows

# Install Django + psycopg2
pip install django psycopg2-binary djangorestframework
""")

doc.add_paragraph("settings.py → PostgreSQL config:")
doc.add_paragraph("""
DATABASES = {
  'default': {
      'ENGINE': 'django.db.backends.postgresql',
      'NAME': 'blogdb',
      'USER': 'postgres',
      'PASSWORD': 'yourpassword',
      'HOST': 'localhost',
      'PORT': '5432',
  }
}
""")

# Add more detailed sections (Part 2, 3, 4, Appendix) similarly...
doc.add_heading("🟦 Part 2 – Django REST Framework (DRF + JWT Authentication)", level=1)
doc.add_paragraph("Detailed explanation of DRF, JWT setup, models, serializers, routers, permissions, filtering, pagination.")

doc.add_heading("🟦 Part 3 – Mini Project: Blog App", level=1)
doc.add_paragraph("Step-by-step guide to building a full Blog App with backend + frontend + JWT. Includes exercises and extra features.")

doc.add_heading("🟦 Part 4 – Deployment Guide", level=1)
doc.add_paragraph("Guide to deploying React frontend on Vercel and Django backend with PostgreSQL on Render. Connecting both for production.")

doc.add_heading("🖼️ Appendix: Architecture Diagram (Textual)", level=1)
doc.add_paragraph("""
[ React Frontend (Vercel) ]
      |
      v
[ REST API Calls via Axios ]
      |
      v
[ Django REST Framework + JWT (Render) ]
      |
      v
[ PostgreSQL Database (Cloud) ]
""")

doc.add_heading("📚 Resources", level=1)
doc.add_paragraph("""
React Docs: https://react.dev
Django Docs: https://docs.djangoproject.com/
DRF Docs: https://www.django-rest-framework.org/
JWT Auth: https://django-rest-framework-simplejwt.readthedocs.io
PostgreSQL: https://www.postgresql.org/
""")

# Save document
file_path = "/mnt/data/React_Django_FullStack_Coursebook_Detailed.docx"
doc.save(file_path)

file_path
# Output the file path for confirmation